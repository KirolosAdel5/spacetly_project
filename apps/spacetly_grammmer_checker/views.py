from rest_framework import generics,status,pagination 
from rest_framework.decorators import api_view,permission_classes
from .models import Document, DocumentImage, DocumentFile
from .serializers import DocumentSerializer, DocumentImageSerializer, DocumentFileSerializer ,RichTextInputSerializer
from rest_framework.response import Response
from .tasks import correct_the_grammar_mistakes
import ast
import uuid
from django.shortcuts import get_object_or_404
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from rest_framework.permissions import IsAuthenticated
import json
from adawat import adaat
from rest_framework.views import APIView
from django.conf import settings
import os
import uuid
from django.http import HttpRequest
from urllib.parse import quote
from django.utils.text import slugify
from .html2word import Html2Word
from .html2pdf import convert_html_to_pdf
from .html2text import convert_html_to_text
class DocumentPagination(pagination.PageNumberPagination):
    page_size = 10
    page_size_query_param = 'page_size'
    max_page_size = 100

class CheckMistakesAPIView(generics.UpdateAPIView):
    def post(self, request, unique_id, *args, **kwargs):
        try:
            document = Document.objects.get(unique_id=unique_id)
        except Document.DoesNotExist:
            document = Document(unique_id=unique_id, created_by=request.user)

        # Update the document fields as needed
        # For example, you can update the content here

        document.save()
       
        # Get the text from the request data
        text = request.data.get('content', '')
        
        # Correct grammar mistakes
        mistakes =  correct_the_grammar_mistakes(text)
        document.mistakes =  mistakes

        # Serialize the updated document data
        serializer = DocumentSerializer(instance=document, data=request.data)
        
        # Validate and save the updated document
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

class CorrrectMistakesAPIView(APIView):
    def post(self, request, unique_id, *args, **kwargs):
        #get or 404 
        document =  get_object_or_404(Document, unique_id=unique_id, created_by=request.user)

       
        # Get the text from the request data
        text = request.data.get('content', '')
        
        # Get the list of mistakes to correct from the request data
        mistakes_to_correct = request.data.get('mistakes_to_correct', [])

        
        # Apply the corrections to the text and update the mistakes in the document
        corrected_text = text
        updated_mistakes = document.mistakes.copy()  # Create a copy to update
        for mistake_group in mistakes_to_correct:
            # Check if the mistake_group is a dictionary
            if isinstance(mistake_group, dict):
                # Iterate over each mistake type and its corrections
                for mistake_type, corrections in mistake_group.items():
                    # Ensure corrections is a dictionary
                    if isinstance(corrections, dict):
                        for uncorrected, corrected in corrections.items():
                            # Check if corrected is a list
                            if isinstance(corrected, list) and corrected:
                                # Apply the first correction from the list
                                corrected_text = corrected_text.replace(uncorrected, corrected[0])
                            elif isinstance(corrected, str):
                                # Apply the correction if it's a string
                                corrected_text = corrected_text.replace(uncorrected, corrected)
                            else:
                                print("Invalid format for correction:", corrected)
                            # Update the mistakes
                            if uncorrected in updated_mistakes.get(mistake_type, {}):
                                del updated_mistakes[mistake_type][uncorrected]
                    else:
                        # Handle the case when corrections is not a dictionary
                        print("Invalid format for corrections:", corrections)
            else:
                # Handle the case when mistake_group is not a dictionary
                print("Invalid format for mistake_group:", mistake_group)

        # Save the corrected text and updated mistakes to the document
        document.corrected_text = corrected_text
        document.mistakes = updated_mistakes
        document.save()        
        
        # Serialize the updated document data
        serializer = DocumentSerializer(instance=document)
        
        # Return the serialized document data
        return Response(serializer.data, status=status.HTTP_200_OK)
    
class DocumentListCreateAPIView(generics.ListCreateAPIView):
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    lookup_field = 'unique_id'  # Specify the lookup field
    pagination_class = DocumentPagination  # Add pagination class
    permission_classes = [IsAuthenticated]
    def get_queryset(self):
        return Document.objects.filter(created_by=self.request.user)
    def create(self, request, *args, **kwargs):
        unique_id = uuid.uuid4()
        return Response({'unique_id': unique_id}, status=status.HTTP_201_CREATED)

class DocumentRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = DocumentSerializer
    lookup_field = 'unique_id'  # Specify the lookup field

    def get_queryset(self):
        return Document.objects.filter(unique_id=self.kwargs['unique_id'])

    def update(self, request, *args, **kwargs):
        partial = kwargs.pop('partial', False)
        instance = self.get_object()
        serializer = self.get_serializer(instance, data=request.data, partial=partial)
        serializer.is_valid(raise_exception=True)
        
        return super().update(request, *args, **kwargs)

class DocumentImageViewSet(generics.ListCreateAPIView):
    serializer_class = DocumentImageSerializer

    def get_queryset(self):
        unique_id = self.kwargs['unique_id']
        return DocumentImage.objects.filter(document__unique_id=unique_id)

    def perform_create(self, serializer):
        unique_id = self.kwargs['unique_id']
        document = get_object_or_404(Document, unique_id=unique_id)
        serializer.save(document=document)

class DocumentImageDetailViewSet(generics.RetrieveUpdateDestroyAPIView):
    queryset = DocumentImage.objects.all()
    serializer_class = DocumentImageSerializer

class DocumentFileViewSet(generics.ListCreateAPIView):
    serializer_class = DocumentFileSerializer

    def get_queryset(self):
        unique_id = self.kwargs['unique_id']
        return DocumentFile.objects.filter(document__unique_id=unique_id)

    def perform_create(self, serializer):
        unique_id = self.kwargs['unique_id']
        document = get_object_or_404(Document, unique_id=unique_id)
        serializer.save(document=document)

class DocumentFileDetailViewSet(generics.RetrieveUpdateDestroyAPIView):
    queryset = DocumentFile.objects.all()
    serializer_class = DocumentFileSerializer

@api_view(['GET'])
def read_document_file(request, document_id, file_id):
    try:
        document = Document.objects.get(unique_id=document_id)
        file = DocumentFile.objects.get(id=file_id, document=document)
    except (Document.DoesNotExist, DocumentFile.DoesNotExist):
        return Response({'error': 'Document or file not found'}, status=status.HTTP_404_NOT_FOUND)

    # Extract text from the file
    if file.file.name.endswith('.docx'):
        text = read_word_file(file.file)
    elif file.file.name.endswith('.pdf'):
        text = read_pdf(file.file)
    else:
        return Response({'error': 'Unsupported file format'}, status=status.HTTP_400_BAD_REQUEST)
    
    # Proofread the extracted text
    proofread_result = text  # Assuming proofread_text() function exists
    return Response({'text': proofread_result})

def read_word_file(uploaded_file):
    full_text = []
    with uploaded_file.open('rb') as file:
        doc = DocxDocument(file)
        for para in doc.paragraphs:
            full_text.append(para.text)
    return '\n'.join(full_text)

def read_pdf(uploaded_file):
    text = ""
    with uploaded_file.open('rb') as file:
        reader = PdfReader(file)
        num_pages = len(reader.pages)
        for page_num in range(num_pages):
            page = reader.pages[page_num]
            text += page.extract_text()
    return text

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def tashkeel_text_APIView(request):
    if request.method == 'POST':
        try:
            text = request.data.get('text')
            lastmark = True if request.data.get('lastmark') == 'True' else False
            if text:
                text_t = adaat.tashkeel_text(text, lastmark)
                return Response({'tashkeel_text': text_t}, status=status.HTTP_200_OK)
            else:
                return Response({'error': 'Text field is required'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    else:
        return Response({'error': 'Method not allowed'}, status=status.HTTP_405_METHOD_NOT_ALLOWED)

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def remove_tashkeel_text_APIView(request):
    if request.method == 'POST':
        try:
            text = request.data.get('text')
            lastmark = True if request.data.get('lastmark') == 'True' else False
            if text:
                text_t = adaat.normalize(text)
                return Response({'normal_text': text_t}, status=status.HTTP_200_OK)
            else:
                return Response({'error': 'Text field is required'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    else:
        return Response({'error': 'Method not allowed'}, status=status.HTTP_405_METHOD_NOT_ALLOWED)

class ExportToPDFView(APIView):
    def get_host_url(self, request: HttpRequest) -> str:
        scheme = request.scheme
        host = request.get_host()
        return f"{scheme}://{host}"

    def post(self, request):
        serializer =  RichTextInputSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        rich_text = serializer.validated_data['rich_text']

        # Save pdf  document to a file
        pdf_filename = f'exported_document_{slugify(uuid.uuid4())}.pdf'
        pdf_path = os.path.join(settings.MEDIA_ROOT, 'pdf_to_export', pdf_filename)
        if convert_html_to_pdf(rich_text, pdf_path):
            # Construct pdf URL
            host_url = self.get_host_url(request)
            pdf_relative_path = os.path.join(settings.MEDIA_URL, 'pdf_to_export', pdf_filename)
            pdf_url = f"{host_url}{quote(pdf_relative_path)}"  # Encoding spaces and special characters
            
            return Response({'pdf_url': pdf_url}, status=status.HTTP_200_OK)
        else :
            return Response({'error': 'Failed to convert HTML to PDF'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ExportToWordView(APIView):
    def get_host_url(self, request: HttpRequest) -> str:
        scheme = request.scheme
        host = request.get_host()
        return f"{scheme}://{host}"

    def post(self, request):
        serializer =  RichTextInputSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        rich_text = serializer.validated_data['rich_text']

        # Save Word document to a file
        word_filename = f'exported_document_{slugify(uuid.uuid4())}.docx'
        word_path = os.path.join(settings.MEDIA_ROOT, 'word_to_export', word_filename)
        try:
            Html2Word(rich_text, word_path)
            # Construct Word URL
            host_url = self.get_host_url(request)
            word_relative_path = os.path.join(settings.MEDIA_URL, 'word_to_export', word_filename)
            word_url = f"{host_url}{quote(word_relative_path)}"  # Encoding spaces and special characters
            
            return Response({'word_url': word_url}, status=status.HTTP_200_OK)
        except Exception:
            return Response({'error': 'Failed to convert HTML to Word'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class ExportToTextView(APIView):
    def get_host_url(self, request: HttpRequest) -> str:
        scheme = request.scheme
        host = request.get_host()
        return f"{scheme}://{host}"

    def post(self, request):
        serializer =  RichTextInputSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        rich_text = serializer.validated_data['rich_text']

        # Save text  document to a file
        txt_filename = f'exported_document_{slugify(uuid.uuid4())}.txt'
        txt_path = os.path.join(settings.MEDIA_ROOT, 'txt_to_export', txt_filename)
        try:
            convert_html_to_text(rich_text, txt_path)
            # Construct text URL
            host_url = self.get_host_url(request)
            txt_relative_path = os.path.join(settings.MEDIA_URL, 'txt_to_export', txt_filename)
            txt_url = f"{host_url}{quote(txt_relative_path)}"  # Encoding spaces and special characters
            
            return Response({'txt_url': txt_url}, status=status.HTTP_200_OK)
        except Exception:
            
            return Response({'error': 'Failed to convert HTML to Text'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)